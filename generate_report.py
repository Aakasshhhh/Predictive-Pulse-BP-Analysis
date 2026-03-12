from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

# Create a Document object
doc = Document()

# Set document margins
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Add title
title = doc.add_heading('🫀 Predictive Pulse – Blood Pressure Prediction System', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_format = title.runs[0].font
title_format.size = Pt(24)
title_format.bold = True
title_format.color.rgb = RGBColor(220, 20, 60)  # Crimson color

# Add subtitle
subtitle = doc.add_paragraph('Comprehensive Project Documentation')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_format = subtitle.runs[0].font
subtitle_format.size = Pt(12)
subtitle_format.italic = True

doc.add_paragraph()  # Spacing

# Add Table of Contents section
doc.add_heading('📋 Table of Contents', level=1)
toc_items = [
    'Project Overview',
    'Features',
    'Technologies Used',
    'Project Structure',
    'Installation',
    'Usage',
    'Data Analysis',
    'Model Training',
    'Web Application',
    'Contributing',
    'License'
]
for item in toc_items:
    doc.add_paragraph(f'• {item}', style='List Bullet')

doc.add_page_break()

# Project Overview Section
doc.add_heading('📌 Project Overview', level=1)
overview_text = """Predictive Pulse is an intelligent machine learning-based web application designed to predict a patient's blood pressure category using various health-related features. The system analyzes patient information and classifies blood pressure into four stages:

• Normal Blood Pressure (Low Risk)
• Stage 1 Hypertension (Moderate Risk)
• Stage 2 Hypertension (High Risk)
• Hypertensive Crisis (Critical Risk)

The application aims to assist in early detection of hypertension risk, helping healthcare professionals and patients make informed decisions about their cardiovascular health."""
doc.add_paragraph(overview_text)

# Features Section
doc.add_heading('🎯 Features', level=1)
features = [
    'Machine Learning Prediction: Utilizes Logistic Regression model trained on patient health data',
    'Web-Based Interface: User-friendly Flask web application for easy access',
    'Comprehensive Health Assessment: Considers multiple health factors including age, gender, family history, symptoms, and vital signs',
    'Risk Classification: Provides clear risk assessment with actionable insights',
    'Data Visualization: Exploratory data analysis with matplotlib and seaborn',
    'Real-time Prediction: Instant results based on user input'
]
for feature in features:
    doc.add_paragraph(feature, style='List Bullet')

# Technologies Section
doc.add_heading('🛠 Technologies Used', level=1)

doc.add_heading('Backend', level=2)
backend_tech = ['Python 3.x', 'Flask - Web framework for the application', 
                'Scikit-learn - Machine learning library for model training and prediction',
                'Joblib - Model serialization and loading']
for tech in backend_tech:
    doc.add_paragraph(tech, style='List Bullet')

doc.add_heading('Data Processing & Analysis', level=2)
data_tech = ['Pandas - Data manipulation and analysis', 'NumPy - Numerical computing',
            'Matplotlib - Data visualization', 'Seaborn - Statistical data visualization']
for tech in data_tech:
    doc.add_paragraph(tech, style='List Bullet')

doc.add_heading('Frontend', level=2)
frontend_tech = ['HTML5 - Web page structure', 'CSS3 - Styling and layout']
for tech in frontend_tech:
    doc.add_paragraph(tech, style='List Bullet')

doc.add_heading('Development Tools', level=2)
dev_tech = ['Jupyter Notebook - Interactive data analysis and visualization']
for tech in dev_tech:
    doc.add_paragraph(tech, style='List Bullet')

doc.add_page_break()

# Project Structure Section
doc.add_heading('📂 Project Structure', level=1)
doc.add_paragraph("""
Predictive-Pulse-BP-Analysis/
│
├── app.py                          # Main Flask application
├── Eda.ipynb                       # Exploratory Data Analysis notebook
├── index.html                      # Static HTML page (if used)
├── README.md                       # Project documentation
│
├── Data/
│   ├── clean_patient_data.csv      # Processed dataset
│   ├── patient_data .csv           # Raw dataset
│   └── src data cleaning.py        # Data cleaning script (duplicate of app.py)
│
├── model/
│   └── modelbuild.py               # Model training script
│
├── static/
│   └── style.css                   # CSS styling for the web app
│
└── templates/
    └── index.html                  # Main web page template
""", style='Normal')

# Installation Section
doc.add_heading('⚙️ Installation', level=1)

doc.add_heading('Prerequisites', level=2)
doc.add_paragraph('Python 3.7 or higher', style='List Bullet')
doc.add_paragraph('pip package manager', style='List Bullet')

doc.add_heading('Step-by-Step Installation', level=2)
doc.add_paragraph('Clone the repository', style='List Number')
doc.add_paragraph('git clone https://github.com/your-username/Predictive-Pulse-BP-Analysis.git', style='Normal')
doc.add_paragraph('cd Predictive-Pulse-BP-Analysis', style='Normal')

doc.add_paragraph('Create a virtual environment (recommended)', style='List Number')
doc.add_paragraph('python -m venv venv', style='Normal')
doc.add_paragraph('source venv/bin/activate  # On Windows: venv\\Scripts\\activate', style='Normal')

doc.add_paragraph('Install required packages', style='List Number')
doc.add_paragraph('pip install -r requirements.txt', style='Normal')
doc.add_paragraph('If requirements.txt does not exist, install the following packages manually:', style='Normal')
doc.add_paragraph('pip install flask scikit-learn pandas numpy matplotlib seaborn joblib', style='Normal')

doc.add_paragraph('Ensure model file is present', style='List Number')
doc.add_paragraph('The trained model logreg_model.pkl should be in the root directory', style='List Bullet')
doc.add_paragraph('If missing, run the model training script: python model/modelbuild.py', style='List Bullet')

doc.add_page_break()

# Usage Section
doc.add_heading('🚀 Usage', level=1)

doc.add_heading('Running the Web Application', level=2)
doc.add_paragraph('Start the Flask application', style='List Number')
doc.add_paragraph('python app.py', style='Normal')

doc.add_paragraph('Open your web browser and navigate to', style='List Number')
doc.add_paragraph('http://127.0.0.1:5000/', style='Normal')

doc.add_paragraph('Enter patient information in the web form including:', style='List Number')
patient_info = ['Age', 'Gender', 'Family history of hypertension', 'Current medication status',
               'Symptoms (breath shortness, vision changes, nose bleeding)',
               'Blood pressure readings (systolic/diastolic)', 'Lifestyle factors (diet, medical care)']
for info in patient_info:
    doc.add_paragraph(info, style='List Bullet')

doc.add_paragraph('Get instant prediction of blood pressure category and risk level', style='List Number')

doc.add_heading('Running Data Analysis', level=2)
doc.add_paragraph('Open the Jupyter notebook', style='List Number')
doc.add_paragraph('jupyter notebook Eda.ipynb', style='Normal')

doc.add_paragraph('Execute cells to perform exploratory data analysis and visualize the dataset', style='List Number')

# Data Analysis Section
doc.add_heading('📊 Data Analysis', level=1)
doc.add_paragraph('The project includes comprehensive exploratory data analysis in Eda.ipynb:')

data_analysis_points = [
    'Data Loading: Import and examine the clean patient dataset',
    'Statistical Summary: Descriptive statistics of all features',
    'Data Visualization: Distribution plots, correlation analysis, categorical feature analysis',
    'Data Quality Assessment: Check for missing values, duplicates, and outliers'
]
for point in data_analysis_points:
    doc.add_paragraph(point, style='List Bullet')

doc.add_heading('Dataset Features', level=2)
dataset_features = ['Age', 'Gender', 'Family History', 'Medication Status',
                   'Symptom Severity', 'Respiratory Symptoms', 'Vision Changes',
                   'Nose Bleeding', 'Diagnosis Timeline', 'Blood Pressure Measurements',
                   'Dietary Habits', 'Medical Care Access']
for feature in dataset_features:
    doc.add_paragraph(feature, style='List Bullet')

doc.add_page_break()

# Model Training Section
doc.add_heading('🧠 Model Training', level=1)

doc.add_heading('Model Building Process (model/modelbuild.py)', level=2)
model_process = [
    'Data Loading: Import cleaned patient data',
    'Preprocessing: Handle missing values and data cleaning',
    'Feature Selection: Select relevant health indicators',
    'Train-Test Split: 80-20 split with random state for reproducibility',
    'Model Selection: Logistic Regression with max_iter=1000',
    'Training: Fit model on training data',
    'Evaluation: Accuracy score, Classification report, Confusion matrix',
    'Model Saving: Serialize trained model using joblib'
]
for step in model_process:
    doc.add_paragraph(step, style='List Number')

doc.add_heading('Model Performance', level=2)
doc.add_paragraph('Algorithm: Logistic Regression', style='List Bullet')
doc.add_paragraph('Target Classes: 4 blood pressure categories (0-3)', style='List Bullet')

# Web Application Section
doc.add_heading('🌐 Web Application', level=1)

doc.add_heading('Flask App Structure (app.py)', level=2)
doc.add_paragraph('Routes:', style='Normal')
doc.add_paragraph('/: Home page with prediction form', style='List Bullet')
doc.add_paragraph('/predict: POST endpoint for prediction requests', style='List Bullet')

doc.add_paragraph('Features:', style='Normal')
web_features = ['Form data validation', 'Model prediction integration',
               'Error handling', 'Result rendering with risk assessment']
for feature in web_features:
    doc.add_paragraph(feature, style='List Bullet')

doc.add_heading('User Interface (templates/index.html, static/style.css)', level=2)
ui_features = [
    'Responsive Design: Works on desktop and mobile devices',
    'Intuitive Form: Easy-to-fill health assessment questionnaire',
    'Clear Results: Prominent display of prediction and risk level',
    'Professional Styling: Clean, medical-themed design'
]
for feature in ui_features:
    doc.add_paragraph(feature, style='List Bullet')

# Contributing Section
doc.add_heading('🤝 Contributing', level=1)
doc.add_paragraph('We welcome contributions to improve Predictive Pulse! Here\'s how you can help:')

contributing_steps = [
    'Fork the repository',
    'Create a feature branch: git checkout -b feature-name',
    'Make your changes and test thoroughly',
    'Commit your changes: git commit -am \'Add new feature\'',
    'Push to the branch: git push origin feature-name',
    'Submit a Pull Request'
]
for step in contributing_steps:
    doc.add_paragraph(step, style='List Number')

doc.add_heading('Areas for Improvement', level=2)
improvements = [
    'Add more machine learning algorithms',
    'Implement cross-validation',
    'Add more comprehensive data validation',
    'Improve UI/UX design',
    'Add user authentication',
    'Implement data persistence'
]
for improvement in improvements:
    doc.add_paragraph(improvement, style='List Bullet')

doc.add_page_break()

# Future Improvements Section
doc.add_heading('🚀 Future Improvements', level=1)
future_items = [
    'Add more medical features',
    'Improve model accuracy using advanced algorithms',
    'Deploy the application online',
    'Add patient dashboard'
]
for item in future_items:
    doc.add_paragraph(item, style='List Bullet')

# License Section
doc.add_heading('📄 License', level=1)
doc.add_paragraph('This project is licensed under the MIT License - see the LICENSE file for details.')

# Contact Section
doc.add_heading('📞 Contact', level=1)
doc.add_paragraph('For questions, suggestions, or collaboration opportunities:')
contact_info = [
    'Project Maintainers: Akash Verma, Shivant Kumar Sahu, Himanshu Aryan, Anjali Kumari',
    'GitHub Issues: Create an issue at the repository'
]
for info in contact_info:
    doc.add_paragraph(info, style='List Bullet')

# Disclaimer Section
doc.add_heading('⚠️ Disclaimer', level=1)
disclaimer_text = 'This application is for educational and informational purposes only. It should not replace professional medical advice, diagnosis, or treatment. Always consult with qualified healthcare providers for medical concerns.'
doc.add_paragraph(disclaimer_text)

doc.add_paragraph()  # Spacing

# Save the document
output_path = 'PROJECT_DOCUMENTATION.docx'
doc.save(output_path)
print(f'✅ Project report document created successfully: {output_path}')
